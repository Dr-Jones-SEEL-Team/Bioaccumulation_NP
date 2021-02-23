# -*- coding: utf-8 -*-


vn_dim_analysis=2.0

import os
import numpy as np
import matplotlib.pyplot as plt
import docx
from docx.shared import Pt
from datetime import datetime
import matplotlib.animation as anim
from matplotlib.animation import FuncAnimation

def dim_analysis(c_set,parameter_combos_count,parameter_matrix,dim_param,new_count_number,vn_N2,vn_Main_Code,vn_parameter_matrix_generator,vn_parameter_checker,vn_csv_generator,vn_method_of_lines,vn_RJ,machine_number,internal_export_path):
    report=docx.Document()
    report.add_heading(f'Results from N2 Run #{new_count_number}-{machine_number}',0) 
    now = datetime.now()
    dt_string = now.strftime("%d/%m/%Y %H:%M:%S")
    date_time_line=report.add_paragraph('Date and Time Report Generated:  ')
    date_time_line.add_run(dt_string)
    report.add_paragraph(f'Machine number ran on- {machine_number}')
    report.add_paragraph(f'N2 version: {vn_N2}')
    report.add_paragraph(f'Main Code version: {vn_Main_Code}')
    report.add_paragraph(f'Parameter Matrix Generator version: {vn_parameter_matrix_generator}')
    report.add_paragraph(f'Parameter Checker version: {vn_parameter_checker}')
    report.add_paragraph(f'CSV Generator version: {vn_csv_generator}')
    report.add_paragraph(f'Method of Lines version: {vn_method_of_lines}')
    report.add_paragraph(f'Residual-Jacobian Calculator version: {vn_RJ}')
    #report.add_paragraph(f'Linear Approximator version: {vn_linear_fitting}')
    style=report.styles['Normal']
    font=style.font
    font.name='Arial'
    font.size=Pt(9)  
    
    # %%unpack dimensional parameters
    Do= dim_param[0] #Diffisivity of NP in the supernatant [m^2/s]
    Dmin=dim_param[1] #Diffusivity of NP in biofilm [m^2/s]
    zeta=dim_param[2] #Zeta potential [Volts, -25 mV] 
    H= dim_param[3] #Biofilm thickness [m, 100 microns]
    kf= dim_param[4] #forward rate constant for binding kinetics [(m^3/s)*(m^3/kg)^(n-1)/(sites)]
    ctot= dim_param[5] #total concentration of binding sites [sites/m^3]
    Kp= dim_param[6] #parition coeffeicent for NP across supernatant-biofilm interface [dimensionless]
    co= dim_param[7] #Concentration of nanoparticle in the supernatant [kg/m^3, 100 ug/L]
    kr= dim_param[8] #Reverse bidning rate constant [1/s]
    phim= dim_param[9] #maximum potential in the biofilm [V]
    n= dim_param[10] #hill coeffecient
    to= H**2/(Do-Dmin) #Calculate dimensionless time
    para0=report.add_paragraph(f'Do={Do}    ')
    para0.add_run(f'Dmin={Dmin}    ')
    para0.add_run(f'zeta={zeta}    ')
    para0.add_run(f'H={H}    ')
    para0.add_run(f'kf={kf}')
    para01=report.add_paragraph(f'ctot={ctot}    ')
    para01.add_run(f'Kp={Kp}   ')
    para01.add_run(f'co={co}   ')
    para01.add_run(f'kr={kr}')
    para02=report.add_paragraph(f'phim={phim}   ')
    para02.add_run(f'n={n}   ')
    para02.add_run(f'to={to}')
    para03.add_run(f'Attempts to recreate Miao et al 2015 doi: 10.1007/s11356-014-3952-y')
    
        
    lin_fit=np.zeros((parameter_combos_count,2))#initialize matrix to store linear best fit parameters
    for pc_i in np.arange(0,parameter_combos_count,1): #Begin for loop to plot the different model paramters using MOL 
        # %% Print Parameters for report
        break_paragraph=report.add_paragraph('___________')
        break_paragraph.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
        report.add_heading('Parameter Set %i'%pc_i,1)
        para1=report.add_paragraph(f'Step-size (h) : {parameter_matrix[pc_i,0]}     ')
        para1.add_run(f'Initial time (t1) : {parameter_matrix[pc_i,2]}     ')
        para1.add_run(f'Final time (t2) : {parameter_matrix[pc_i,3]}     ')
        para1.add_run(f'Mesh size (nx) : {parameter_matrix[pc_i,4]}')
        para2=report.add_paragraph(f'Dimensionless ratio of diffusivity (gamma) : {parameter_matrix[pc_i,5]}     ')
        para2.add_run(f'Dimensionless ratio of potential (beta): {parameter_matrix[pc_i,6]}')
        para3=report.add_paragraph(f'Dimensionless forward rate constant (F): {parameter_matrix[pc_i,7]}     ')
        para3.add_run(f'Dimensionless reverse rate constant (R): {parameter_matrix[pc_i,8]}')
        para4=report.add_paragraph(f'Hill coeffecient (n): {parameter_matrix[pc_i,9]}     ')
        para4.add_run(f'Tolerance: {parameter_matrix[pc_i,1]}')
    
    
        # %% Grab Dimensionless Concentration and times from c_set and convert to dimensionless concentrations and times
        ncb=c_set[pc_i][0] #Grab current dimensionless bound concentration data to plot
        cb=ncb*ctot #Convert dimensionless bound concentration vector to dimensional bound concentration vector
        ncu=c_set[pc_i][1] #Grab current dimenseionless unbound concentration data to plot
        cu=ncu*Kp*co #Convert dimensionless bound concentration vector to dimensional bound concentration vector
        nt=c_set[pc_i][6] #Grab dimensionless time-vector for this parameter set for plotting
        t=nt*to #claculate dimensional time-vector
        nt=c_set[pc_i][7] #Grab number of time points for this parameter set for plotting
        nx=c_set[pc_i][8] #Grab the dimensionless position-vector for this parameter set for plotting
        x=nx*H #Calculate dimensionla position vector for plotting 
        
        # %% Find Average Dimensional Bound Concentration Overtime
        average_bound_conc=np.zeros(nt) #Initiate average bound concentration overtime vector
        t2index=np.arange(0,nt) #Create time vector to loop over
        for t2_i in t2index:
            average_bound_conc[t2_i]=np.average(cb[:,t2_i])
            
        # %% Find Average Dimensional Unbound Concentration Overtime
        average_unbound_conc=np.zeros(nt) #Initiate average unbound concentration overtime vector
        for t2_i in t2index:
             average_unbound_conc[t2_i]=np.average(cu[:,t2_i])
             
        # %% Calculate Total Concentration Overtime and and Average
        ct=cu+cb #total concentration 
        average_total_conc=np.zeros(nt) #Initiate average total concentration overtime vector
        t2index=np.arange(0,nt) #Create time vector to loop over
        for t2_i in t2index:
            average_total_conc[t2_i]=np.average(ct[:,t2_i])
            
        # %% Calculating log of equilibrium-normalized concentration overtime (for first-order transfer kinetics fit)
        avg_total_conc_equil= Kp*co+n*ctot/(1+kr*n/kf/(Kp*co)**n) #average total concentration of NP at equilibrium 
        log_norm_conc=np.log(avg_total_conc_equil-average_total_conc) #Calculate the log of equilibrium normalized concentration overtime
        perc_acc_model=average_total_conc/avg_total_conc_equil #Calculate percent of NP accumulated with respect to equilibrium
        
        # %% Linear Fits analysis
        # Find the 99% accumulation time and cutoff perc_acc_model
        cutoff= 0.99 #cutoff percentage to "reach equilibrium" (used to determine where to start fit)
        j=0 #reset counter for time-search loop
        mod_cutoff=0 #Initialize mod-cutoff
        for t_i in t:
            if perc_acc_model[j]>cutoff:
                mod_cutoff=j #Index neccesary for model to reach evaluated percent accumulated
                break
            j=j+1 #Update counter in time search loop
        if mod_cutoff == 0 : print(f'did not reach {cutoff} of equilibrium')
        t_cutoff=t[:mod_cutoff]
        log_norm_conc_cutoff=log_norm_conc[:mod_cutoff]

        # %% Run linear fit
        [m,b]=np.polyfit(t_cutoff,log_norm_conc_cutoff,1) #Find linear fit for plot
        lin_fit[pc_i,0]=m
        lin_fit[pc_i,1]=b #Store fit into matrix
        
        # %% Use parameters to generate best-fit data
        fit_ct=np.zeros(len(t))#Initialize concentration vector
        fit_log_norm_conc=np.zeros(len(t)) #Initialize longnormal concentration vactor
        count=0 #Begin counter
        for t_i in t:
            fit_ct[count]=avg_total_conc_equil*(1-np.exp(m*t_i))
            fit_log_norm_conc[count]=m*t_i+b
            count=count+1
        
        # %%Find relelvant maximums and minimums
        upper_0 = np.amax(cu)*1.1 #Upper bound on Unbound Concentration
        upper_1 = np.amax(cb)*1.1 #Upper Bound on Average Unbound Concentration
        upper_2 = np.amax(ct)*1.1 #Upper bound on total NP concentration
        upper_3 = np.amax(average_unbound_conc)*1.1 #Upper Bound on Change in Average Concentration
        upper_4 = np.amax(average_bound_conc)*1.1 #Upper bound on Bound Concentration
        upper_5 = np.amax(average_total_conc)*1.1 #Upper Bound on Average total Concentration
        upper_6 = np.amax(log_norm_conc)*1.1 #Upper Bound on log-eq-norm conc plot (should be 0.9 if max is negative, 1.1 if max is positive)
        lower_6 = np.amin(log_norm_conc)*0.9 #Lower bound on log-eq-norm conc plot (should be 0.9 if min is positive, 1.1 if min is negative)
        upper_7 = np.amax(x) #Maximum position in domain
        upper_8 = np.amax(fit_ct)*1.1 #Upper bound on fit average total concentration overtime
        upper_9 = np.amax(fit_log_norm_conc)*1.1 #Upper Bound on 1st order fit of log-eq-norm conc plot (should be 0.9 if max is negative, 1.1 if max is positive)
        lower_9 = np.amin(fit_log_norm_conc)*0.9 #Lower Bound on 1st order fit of log-eq-norm conc plot (should be 0.9 if min is positive, 1.1 if min is negative)
        
        # %%Unbound Concentration vs position plot
        #tindex_u=np.array([0,5,10,25,50,75,100,125,150,200,250]) for masnual control over timepoints plotted
        tp_u=10 #number of time points to plot
        plt.figure(8*pc_i+0)
        
        #Linear discretization of plotted timepionts
        space_u=int((nt-1)/tp_u) #Linear discreitzation of timepoints
        tindex_u=np.arange(0,nt,space_u) #Linear discreitization of timepoints
        for i_u in tindex_u:
            cc_u=cu[:,i_u]
            ti_u=round(t[i_u],5)
            plt.plot(x,cc_u,label='t={}'.format(ti_u))
        """
        #Logarthmic discreitzation of plotted timepoints
        lognt_u=np.log10(nt) #Logarthmic timepoints
        logspace_u=round((lognt_u)/tp_u,10) #Logarthmic timepoints
        logtindex_u=np.arange(0,lognt_u,logspace_u) #Logarthmic timepoints
        plt.figure(8*pc_i+0)
        for logi_u in logtindex_u:
            i_u=int(10**logi_u)
            cc_u=cu[:,i_u]
            ti_u=round(t[i_u],5)
            plt.plot(x,cc_u,label=f't={ti_u}')
            """
        
        plt.xlim(left=0,right=upper_7)
        plt.ylim(bottom=0,top=upper_0)
        plt.xlabel('Position [m]',fontsize=14)
        plt.ylabel('Concentration [kg/m^3]',fontsize=14)
        plt.title('Unbound Concentration plot',fontsize=16)
        plt.xticks(fontsize=12)
        plt.yticks(fontsize=12)
        plt.legend(loc=(0.1,0.1))
        unbound_filename_partial=f'Unboundplot{pc_i}.png'
        unbound_filename_full=os.path.join(internal_export_path,unbound_filename_partial)
        plt.savefig(unbound_filename_full)
        plt.close()
       
        
        # %%Bound Concentration vs position plot
        tp_b=10 #number of time points to plot
        plt.figure(8*pc_i+1)
        
        #Linear discretization of plotted timepionts
        space_b=int((nt-1)/tp_b) #Linear discreitzation of timepoints
        tindex_b=np.arange(0,nt,space_b) #Linear discreitzation of timepoints
        for i_b in tindex_b:
            cc_b=cb[:,i_b]
            ti_b=round(t[i_b],5)
            plt.plot(x,cc_b,label='t={}'.format(ti_b))
        
        """
        #Logarthmic discreitzation of plotted timepoints
        lognt_b=np.log10(nt) #Logarthmic timepoints
        logspace_b=round((lognt_b)/tp_b,10) #Logarthmic timepoints
        logtindex_b=np.arange(0,lognt_b,logspace_b) #Logarthmic timepoints
        plt.figure(8*pc_i+1)
        for logi_b in logtindex_b:
            i_b=int(10**logi_b)
            cc_b=cb[:,i_b]
            ti_b=round(t[i_b],5)
            plt.plot(x,cc_b,label=f't={ti_b}')
            """
            
        plt.xlim(left=0,right=upper_7)
        plt.ylim(bottom=0,top=upper_1)
        plt.xlabel('Position [m]',fontsize=14)
        plt.ylabel('Concentration [kg/m^3]',fontsize=14)
        plt.title('Bound Concentration plot',fontsize=16)
        plt.xticks(fontsize=12)
        plt.yticks(fontsize=12)
        plt.legend(loc=(0.1,0.1))
        bound_filename_partial=f'Boundplot{pc_i}.png'
        bound_filename_full=os.path.join(internal_export_path,bound_filename_partial)
        plt.savefig(bound_filename_full)
        plt.close()
        pics_paragraph1=report.add_paragraph()
        pic1=pics_paragraph1.add_run()
        pic1.add_picture(unbound_filename_full, width=docx.shared.Inches(3))
        pic2=pics_paragraph1.add_run()
        pic2.add_picture(bound_filename_full, width=docx.shared.Inches(3))
        
        # %% Total concentration plot vs position
        tp_t=10 #number of time points to plot
        plt.figure(8*pc_i+2)
        
        #Linear discretization of plotted timepionts
        space_b=int((nt-1)/tp_b) #Linear discreitzation of timepoints
        tindex_b=np.arange(0,nt,space_b) #Linear discreitzation of timepoints
        for i_b in tindex_b:
            cc_b=cb[:,i_b]
            ti_b=round(t[i_b],5)
            plt.plot(x,cc_b,label='t={}'.format(ti_b))
        """
        
        #Logarthmic discreitzation of plotted timepoints
        lognt_t=np.log10(nt) #Logarthmic timepoints
        logspace_t=round((lognt_t)/tp_t,10) #Logarthmic timepoints
        logtindex_t=np.arange(0,lognt_t,logspace_t) #Logarthmic timepoints
        plt.figure(8*pc_i+2)
        for logi_t in logtindex_t:
            i_t=int(10**logi_t)
            cc_t=ct[:,i_t]
            ti_t=round(t[i_t],5)
            plt.plot(x,cc_b,label=f't={ti_t}')
            """
        plt.xlim(left=0,right=upper_7)
        plt.ylim(bottom=0,top=upper_2)
        plt.xlabel('Position [m]',fontsize=14)
        plt.ylabel('Concentration [kg/m^3]',fontsize=14)
        plt.title('Total NP Concentration plot',fontsize=16)
        plt.xticks(fontsize=12)
        plt.yticks(fontsize=12)
        plt.legend(loc=(0.1,0.1))
        total_filename_partial=f'Totalplot{pc_i}.png'
        total_filename_full=os.path.join(internal_export_path,total_filename_partial)
        plt.savefig(total_filename_full)
        plt.close()

        
        # %% Unbound NP Average Concetration Overtime
        plt.figure(8*pc_i+3)
        plt.plot(t,average_unbound_conc)
        plt.xlim(left=parameter_matrix[pc_i,2]*to,right=parameter_matrix[pc_i,3]*to)
        #plt.xlim(left=0,right=0.0005) #Manual Override of automatic x-axis limits
        plt.ylim(bottom=0,top=upper_3)
        plt.xlabel('Time [s]',fontsize=14)
        plt.ylabel('Concentration [kg/m^3]',fontsize=14)
        plt.title('Average Unbound Concentration plot',fontsize=16)
        plt.xticks(fontsize=12)
        plt.yticks(fontsize=12)
        avgunbound_filename_partial=f'Avg Unboundplot{pc_i}.png'
        avgunbound_filename_full=os.path.join(internal_export_path,avgunbound_filename_partial)
        plt.savefig(avgunbound_filename_full)
        plt.close()
        pics_paragraph2=report.add_paragraph()
        pic3=pics_paragraph2.add_run()
        pic3.add_picture(total_filename_full, width=docx.shared.Inches(3))
        pic4=pics_paragraph2.add_run()
        pic4.add_picture(avgunbound_filename_full, width=docx.shared.Inches(3))
        
        """# %%Unbound NP Change in Concentration vs Concentration
        plt.figure(7*pc_i+3)
        plt.plot(average_conc_overtime,change_in_concentration)
        plt.xlim(left=0,right=upper_2)
        plt.ylim(bottom=0,top=upper_3)
        plt.xlabel('Concentration',fontsize=14)
        plt.ylabel('Dimensionless Change in Concentration',fontsize=14)
        plt.title('Unbound dC vs C plot',fontsize=16)
        plt.xticks(fontsize=12)
        plt.yticks(fontsize=12)
        dCunbound_filename_partial=f'dC Unboundplot{pc_i}.png'
        dCunbound_filename_full=os.path.join(internal_export_path,dCunbound_filename_partial)
        plt.savefig(dCunbound_filename_full)
        plt.close()
        pics_paragraph2=report.add_paragraph()
        pic3=pics_paragraph2.add_run()
        pic3.add_picture(avgunbound_filename_full, width=docx.shared.Inches(3))
        pic4=pics_paragraph2.add_run()
        pic4.add_picture(dCunbound_filename_full, width=docx.shared.Inches(3))
        """
        # %% Bound NP Average Concetration Overtime
        plt.figure(8*pc_i+4)
        plt.plot(t,average_bound_conc)
        plt.xlim(left=parameter_matrix[pc_i,2]*to,right=parameter_matrix[pc_i,3]*to)
        #plt.xlim(left=0,right=0.0005) #Manual Override of automatic x-axis limits
        plt.ylim(bottom=0,top=upper_4)
        plt.xlabel('Time [s]',fontsize=14)
        plt.ylabel('Concentration [kg/m^3]',fontsize=14)
        plt.title('Average Bound Concentration plot',fontsize=16)
        plt.xticks(fontsize=12)
        plt.yticks(fontsize=12)
        avgbound_filename_partial=f'Avg Boundplot{pc_i}.png'
        avgbound_filename_full=os.path.join(internal_export_path,avgbound_filename_partial)
        plt.savefig(avgbound_filename_full)
        plt.close()
        
        # %% Total NP Average Concetration Overtime
        plt.figure(8*pc_i+5)
        plt.plot(t,average_total_conc)
        plt.xlim(left=parameter_matrix[pc_i,2]*to,right=parameter_matrix[pc_i,3]*to)
        #plt.xlim(left=0,right=0.0005) #Manual Override of automatic x-axis limits
        plt.ylim(bottom=0,top=upper_5)
        plt.xlabel('Time [s]',fontsize=14)
        plt.ylabel('Concentration [kg/m^3]',fontsize=14)
        plt.title('Average Total Concentration plot',fontsize=16)
        plt.xticks(fontsize=12)
        plt.yticks(fontsize=12)
        avgtotal_filename_partial=f'Avg totalplot{pc_i}.png'
        avgtotal_filename_full=os.path.join(internal_export_path,avgtotal_filename_partial)
        plt.savefig(avgtotal_filename_full)
        plt.close()
        pics_paragraph3=report.add_paragraph()
        pic5=pics_paragraph3.add_run()
        pic5.add_picture(avgbound_filename_full, width=docx.shared.Inches(3))
        pic6=pics_paragraph3.add_run()
        pic6.add_picture(avgtotal_filename_full, width=docx.shared.Inches(3))
        
        # %% Log Equilibrium concentration vs time plot (first-order fit)   
        plt.figure(8*pc_i+6)
        plt.plot(t,log_norm_conc,label='Model Results')
        plt.plot(t,fit_log_norm_conc,label='First-Order Approximation')
        if upper_9>upper_6:
            uplimit=upper_9
        else: uplimit=upper_6
        if lower_9<lower_6:
            lowlimit=lower_9
        else: lowlimit=lower_6
        plt.xlim(left=parameter_matrix[pc_i,2]*to,right=parameter_matrix[pc_i,3]*to)  
        plt.ylim(bottom=lowlimit,top=uplimit)
        plt.xlabel('Time',fontsize=14)
        plt.ylabel('log(Normalized Average Concentration',fontsize=14)
        plt.title('Model vs First-order Mass Transfer',fontsize=16)
        plt.xticks(fontsize=12)
        plt.yticks(fontsize=12)
        plt.legend(loc=(0.5,0.75))
        log_filename_partial=f'Logplot{pc_i}.png'
        log_filename_full=os.path.join(internal_export_path,log_filename_partial)
        plt.savefig(log_filename_full)
        plt.close()
        
        # %% Average Total Concentration vs First-Order Approximation
        plt.figure(8*pc_i+7)
        plt.plot(t,average_total_conc,label='Model Results')
        plt.plot(t,fit_ct,label='First-Order Approximation')
        if upper_8>upper_5:
            uplimit2=upper_8
        else: uplimit2=upper_5
        plt.xlim(left=parameter_matrix[pc_i,2]*to,right=parameter_matrix[pc_i,3]*to) 
        plt.ylim(bottom=0,top=uplimit2)
        plt.xlabel('Time [s]',fontsize=14)
        plt.ylabel('Average Total Concentration [kg/m^3]',fontsize=14)
        plt.title('Model vs First-order Mass Transfer',fontsize=16)
        plt.xticks(fontsize=12)
        plt.yticks(fontsize=12)
        plt.legend(loc=(0.5,0.3))
        plt.text(0.4,0.5,f'm={m}, b={b}')
        linear_filename_partial=f'Linearplot{pc_i}.png'
        linear_filename_full=os.path.join(internal_export_path,linear_filename_partial)
        plt.savefig(linear_filename_full)
        plt.close()
        pics_paragraph4=report.add_paragraph()
        pic7=pics_paragraph4.add_run()
        pic7.add_picture(log_filename_full, width=docx.shared.Inches(3))
        pic8=pics_paragraph3.add_run()
        pic8.add_picture(linear_filename_full, width=docx.shared.Inches(3))
        
        
        """# %%Total NP Change in Concentration vs Concentration
        plt.figure(7*pc_i+5)
        plt.plot(taverage_conc_overtime,tchange_in_concentration)
        plt.xlim(left=0,right=upper_5)
        plt.ylim(bottom=0,top=upper_6)
        plt.xlabel('Concentration',fontsize=14)
        plt.ylabel('Dimensionless Change in Concentration',fontsize=14)
        plt.title('Total dC vs C plot',fontsize=16)
        plt.xticks(fontsize=12)
        plt.yticks(fontsize=12)
        dCtotalNP_filename_partial=f'dC totolNPplot{pc_i}.png'
        dCtotalNP_filename_full=os.path.join(internal_export_path,dCtotalNP_filename_partial)
        plt.savefig(dCtotalNP_filename_full)
        plt.close()
        pics_paragraph3=report.add_paragraph()
        pic5=pics_paragraph3.add_run()
        pic5.add_picture(avgtotalNP_filename_full, width=docx.shared.Inches(3))
        pic6=pics_paragraph3.add_run()
        pic6.add_picture(dCtotalNP_filename_full, width=docx.shared.Inches(3))
        """
        
        
        """# %%Logarithms of Total NP Change in Concentration vs Concentration
        plt.figure(7*pc_i+6)
        plt.plot(logtavg_conc_overtime,logtchange_conc)
        plt.xlim(left=lower_7,right=upper_7)
        plt.ylim(bottom=lower_8,top=upper_8) 
        plt.xlabel('Log of Dimensionless Concentration',fontsize=14)
        plt.ylabel('Log of Dimensionless Change in Concentration',fontsize=14)
        plt.title('Total log(dC) vs log(C) plot',fontsize=16)
        plt.xticks(fontsize=12)
        plt.yticks(fontsize=12)
        logdCtotalNP_filename_partial=f'logdC totolNPplot{pc_i}.png'
        logdCtotalNP_filename_full=os.path.join(internal_export_path,logdCtotalNP_filename_partial)
        plt.savefig(logdCtotalNP_filename_full)
        plt.close()
        pics_paragraph4=report.add_paragraph()
        pic7=pics_paragraph4.add_run()
        pic7.add_picture(logdCtotalNP_filename_full, width=docx.shared.Inches(3))
        """
        """
        # %%Unbound Concentration Animation
        unbound_anim_fig=plt.figure()
        unbound_anim_plot=plt.plot([])
        unbound_anim_holder=unbound_anim_plot[0]
        plt.xlim(left=0,right=1)
        plt.ylim(bottom=0,top=upper_1)
        plt.xlabel('Position',fontsize=14)
        plt.ylabel('Dimensionless Concentration',fontsize=14)
        plt.title('Dimensionless Unbound Concentration',fontsize=16)
        plt.xticks(fontsize=12)
        plt.yticks(fontsize=12)
        tp_u_anim=100 #number of time points to plot for the animation for unbound concntration
        space_u_anim=int((nt-1)/tp_u_anim) #space between timepoints    
        def unbound_animate(frame):
            #update plot
            c_u_plot=cu[:,frame*space_u_anim]
            unbound_anim_holder.set_data((x,c_u_plot))
        unbound_anim=anim.FuncAnimation(unbound_anim_fig,unbound_animate,frames=tp_u_anim,interval=100)
        unbound_anim_filename_partial=f'unboun_anim{pc_i}.gif'
        unbound_anim_filename_full=os.path.join(internal_export_path,unbound_anim_filename_partial)
        unbound_anim.save(unbound_anim_filename_full)
        #Only need these lines if log plot is turned off
        pics_paragraph3=report.add_paragraph()
        pic5=pics_paragraph3.add_run()
        #End of possibly neccesary lines
        pic5.add_picture(unbound_anim_filename_full, width=docx.shared.Inches(3))
        plt.close()
        
        """
        
        """# %%Total NP Change in Concentration vs Concentration Animation
        dCvC_anim_fig=plt.figure()
        dCvC_anim_plot=plt.plot([])
        dCvC_anim_holder=dCvC_anim_plot[0]
        plt.xlim(left=0,right=upper_5)
        plt.ylim(bottom=0,top=upper_6)
        plt.xlabel('Concentration',fontsize=14)
        plt.ylabel('Dimensionless Change in Concentration',fontsize=14)
        plt.title('Total dC vs C plot',fontsize=16)
        plt.xticks(fontsize=12)
        plt.yticks(fontsize=12)
        tp_dCvC_anim=100 #number of time points to plot for the animation for unbound concntration
        space_dCvC_anim=int((nt-1)/tp_dCvC_anim) #space between timepoints    
        def dCvC_animate(frame):
            #update plot
            dCvC_anim_holder.set_data((taverage_conc_overtime[0:frame*space_dCvC_anim],tchange_in_concentration[0:frame*space_dCvC_anim]))
        dCvC_anim=anim.FuncAnimation(dCvC_anim_fig,dCvC_animate,frames=tp_dCvC_anim,interval=100)
        dCvC_anim_filename_partial=f'dCvC_anim{pc_i}.gif'
        dCvC_anim_filename_full=os.path.join(internal_export_path,dCvC_anim_filename_partial)
        dCvC_anim.save(dCvC_anim_filename_full)
        #pics_paragraph5=report.add_paragraph() commented out when no log plot
        #pic8=pics_paragraph5.add_run()
        pic7.add_picture(dCvC_anim_filename_full, width=docx.shared.Inches(3))
        """
        """# %%Total NP Concentration vs Time Animation
        totCvt_anim_fig=plt.figure()
        totCvt_anim_plot=plt.plot([])
        totCvt_anim_holder=totCvt_anim_plot[0]
        plt.xlim(left=parameter_matrix[pc_i,2],right=parameter_matrix[pc_i,3])
        plt.ylim(bottom=0,top=upper_5)
        plt.xlabel('Time',fontsize=14)
        plt.ylabel('Dimensionless Concentration',fontsize=14)
        plt.title('Average Dimensionless Total Concentration plot',fontsize=16)
        plt.xticks(fontsize=12)
        plt.yticks(fontsize=12)
        tp_totCvt_anim=100 #number of time points to plot for the animation for unbound concntration
        space_totCvt_anim=int((nt-1)/tp_totCvt_anim) #space between timepoints    
        def totCvt_animate(frame):
            #update plot
            totCvt_anim_holder.set_data((t[0:frame*space_totCvt_anim],taverage_conc_overtime[0:frame*space_totCvt_anim]))
        totCvt_anim=anim.FuncAnimation(totCvt_anim_fig,totCvt_animate,frames=tp_totCvt_anim,interval=100)
        totCvt_anim_filename_partial=f'totCvt_anim{pc_i}.gif'
        totCvt_anim_filename_full=os.path.join(internal_export_path,totCvt_anim_filename_partial)
        totCvt_anim.save(totCvt_anim_filename_full)
        #pics_paragraph5=report.add_paragraph() #Added when log plot off
        pic6=pics_paragraph3.add_run() #Added when log plot off
        pic6.add_picture(totCvt_anim_filename_full, width=docx.shared.Inches(3))
        plt.close()
        
        # %% Plot Approximation for Total NP conc Overtime
        linear_filename_partial=f'Linearplot{pc_i}.png'
        linear_filename_full=os.path.join(internal_export_path,linear_filename_partial)
        #pics_paragraph6=report.add_paragraph() commented out  when log plots out
        #pic9=pics_paragraph6.add_run() commented out  when log plots out
        pics_paragraph4=report.add_paragraph()
        pic7=pics_paragraph4.add_run()
        pic7.add_picture(linear_filename_full,width=docx.shared.Inches(3))
        log_filename_partial=f'Logplot{pc_i}.png'
        log_filename_full=os.path.join(internal_export_path,log_filename_partial)
        pic8=pics_paragraph3.add_run() #Added when log plot off
        pic8.add_picture(log_filename_full,width=docx.shared.Inches(3))
        
        # %% Add Table for Fit
        perc_acc_table=perc_acc_matrix[pc_i][0]
        [table_rows,table_columns]=perc_acc_table.shape
        table1=report.add_table(rows=table_rows+1, cols=table_columns)
        row=table1.rows[0]
        row.cells[0].text='Percent Accumulated'
        row.cells[1].text='Time for Model'
        row.cells[2].text='Time for Approximation'
        row.cells[3].text='Percent Error'
        i_v = np.arange(0,table_rows,1) #index for rows of percent accumulation table
        j_v = np.arange(0,table_columns,1) #inde for columns of percent accumulation table
        for i in i_v:
            for j in j_v:
                cell=table1.cell(i+1,j)
                cell.text=str(perc_acc_table[i,j])
        """
        # %%
    return report

# %%
"""
Created on Sat Feb 13 12:49:16 2021

@author: joshuaprince


Purpose: Script to convert results to dimensional form, and generate report on those results

Version 2.0

Script was first developed for N2.2



"""