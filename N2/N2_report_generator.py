#!/usr/bin/env python3
# -*- coding: utf-8 -*-

vn_report_generator=1.6

import os
import numpy as np
import matplotlib.pyplot as plt
import docx
from docx.shared import Pt
from datetime import datetime
import matplotlib.animation as anim
from matplotlib.animation import FuncAnimation

def plot_generator(c_set,parameter_combos_count,parameter_matrix,new_count_number,vn_N2,vn_Main_Code,vn_parameter_matrix_generator,vn_parameter_checker,vn_csv_generator,vn_method_of_lines,vn_RJ):
    """Static Plotting (Exported to Word Document)"""
    internal_export_path='/Users/joshuaprince/Northeastern University/Jones SEEL Team - Bioremediation of Nanoparticles/Modelling Work/Model Results/N2/Internal Exports' #Indirect Export path for Files, used for outputs which only get using internally
    report=docx.Document()
    report.add_heading(f'Results from N2 Run #{new_count_number}',0)
    now = datetime.now()
    dt_string = now.strftime("%d/%m/%Y %H:%M:%S")
    date_time_line=report.add_paragraph('Date and Time Report Generated:  ')
    date_time_line.add_run(dt_string)
    report.add_paragraph(f'N2 version: {vn_N2}')
    report.add_paragraph(f'Main Code version: {vn_Main_Code}')
    report.add_paragraph(f'Parameter Matrix Generator version: {vn_parameter_matrix_generator}')
    report.add_paragraph(f'Parameter Checker version: {vn_parameter_checker}')
    report.add_paragraph(f'CSV Generator version: {vn_csv_generator}')
    report.add_paragraph(f'Method of Lines version: {vn_method_of_lines}')
    report.add_paragraph(f'Residual-Jacobian Calculator version: {vn_RJ}')
    report.add_paragraph(f'Report Generator version: {vn_report_generator}')
    style=report.styles['Normal']
    font=style.font
    font.name='Arial'
    font.size=Pt(9)  
        
     
    for pc_i in np.arange(0,parameter_combos_count,1): #Begin for loop to plot the different model paramters using MOL 
        cb=c_set[pc_i][0] #Grab current bound concentration data to plot
        cu=c_set[pc_i][1] #Grab current unbound concentration data to plot
        average_conc_overtime=c_set[pc_i][2] #Grab current average concentration data to plot (Unbound NP)
        change_in_concentration=c_set[pc_i][3] #Grab current change in concentration data to plot (Unbound NP)
        taverage_conc_overtime=c_set[pc_i][4] #Grab current change in concentration data to plot (total NP)
        tchange_in_concentration=c_set[pc_i][5] #Grab current change in concentration data to plot (Unbound NP)
        npaverage_conc_overtime=np.array(taverage_conc_overtime) #Convert taverage_conc_overtime array into np.array
        logtavg_conc_overtime=np.log10(npaverage_conc_overtime)  #Logarithm taverage change in concentration overtime
        nptchange_conc=np.log10(tchange_in_concentration) #total change in concentration array conerted to numpy array
        logtchange_conc=np.array(nptchange_conc) #Convert average_conc_overtime array into np.array
        t=c_set[pc_i][6] #Grab time-vector for this parameter set for plotting
        nt=c_set[pc_i][7] #Grab number of time points for this parameter set for plotting
        x=c_set[pc_i][8] #Grab the position-vector for this parameter set for plotting
        break_paragraph=report.add_paragraph('___________')
        break_paragraph.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
        report.add_heading('Parameter Set %i'%pc_i,1)
        para1=report.add_paragraph(f'Step-size (h) : {parameter_matrix[pc_i,0]}     ')
        para1.add_run(f'Initial time (t1) : {parameter_matrix[pc_i,2]}     ')
        para1.add_run(f'Final time (t2) : {parameter_matrix[pc_i,3]}     ')
        para1.add_run(f'Mesh size (nx) : {parameter_matrix[pc_i,4]}')
        para2=report.add_paragraph(f'Dimensionless ratio of diffusivity (gamma) : {parameter_matrix[pc_i,5]}     ')
        para2.add_run(f'Dimensionless ratio of potential (beta): {parameter_matrix[pc_i,6]}')
        para3=report.add_paragraph(f'Dimensionless forward rate constant (F): {parameter_matrix[pc_i,7]}     ')
        para3.add_run(f'Dimensionless reverse rate constant (R): {parameter_matrix[pc_i,8]}')
        para4=report.add_paragraph(f'Hill coeffecient (n): {parameter_matrix[pc_i,9]}     ')
        para4.add_run(f'Tolerance: {parameter_matrix[pc_i,1]}')
        
        #Find relelvant maximums and minimums
        upper_1 = np.amax(cu)*1.1 #Upper bound on Unbound Concentration
        upper_2 = np.amax(average_conc_overtime)*1.1 #Upper Bound on Average Unbound Concentration
        upper_3 = np.amax(change_in_concentration)*1.1 #Upper Bound on Change in Average Concentration
        upper_4 = np.amax(cb)*1.1 #Upper bound on Bound Concentration
        upper_5 = np.amax(taverage_conc_overtime)*1.1 #Upper Bound on Average total Concentration
        upper_6 = np.amax(tchange_in_concentration)*1.1 #Upper Bound on total Change in Average Concentration
        upper_7 = np.amax(logtavg_conc_overtime)*1.1 #Upper bound on log of total average NP conc
        upper_8 = np.amax(logtchange_conc)*1.1 #Upper bound on log of total change in conc
        lower_7 = np.amin(logtavg_conc_overtime)*1.1 #Lower bound on log of total average NP conc
        lower_8 = np.amin(logtavg_conc_overtime)*1.1 #Upper bound on log of total change in conc
        
        #Unbound
        #tindex_u=np.array([0,5,10,25,50,75,100,125,150,200,250]) for masnual control over timepoints plotted
        tp_u=10 #number of time points to plot
        """
        #Linear discretization of plotted timepionts
        #space_u=int((nt-1)/tp_u) #Linear discreitzation of timepoints
        #tindex_b=np.arange(0,nt,space_u) #Linear discreitization of timepoints
        for i_u in tindex_u:
            cc_u=cu[:,i_u]
            ti_u=round(t[i_u],5)
            plt.plot(x,cc_u,label='t={}'.format(ti_u))
        """
        #Logarthmic discreitzation of plotted timepoints
        lognt_u=np.log10(nt) #Logarthmic timepoints
        logspace_u=round((lognt_u)/tp_u,3) #Logarthmic timepoints
        logtindex_u=np.arange(0,lognt_u,logspace_u) #Logarthmic timepoints
        plt.figure(7*pc_i+0)
        for logi_u in logtindex_u:
            i_u=int(10**logi_u)
            cc_u=cu[:,i_u]
            ti_u=round(t[i_u],5)
            plt.plot(x,cc_u,label='t={}'.format(ti_u))
        
        plt.xlim(left=0,right=1)
        plt.ylim(bottom=0,top=upper_1)
        plt.xlabel('Position',fontsize=14)
        plt.ylabel('Dimensionless Concentration',fontsize=14)
        plt.title('Dimensionless Unbound Concentration plot',fontsize=16)
        plt.xticks(fontsize=12)
        plt.yticks(fontsize=12)
        plt.legend(loc=(0.1,0.1))
        unbound_filename_partial=f'Unboundplot{pc_i}.png'
        unbound_filename_full=os.path.join(internal_export_path,unbound_filename_partial)
        plt.savefig(unbound_filename_full)
        plt.close()
       
        
        #Bound
        tp_b=10 #number of time points to plot
        """
        #Linear discretization of plotted timepionts
        #space_b=int((nt-1)/tp_b) #Linear discreitzation of timepoints
        #tindex_b=np.arange(0,nt,space_b) #Linear discreitzation of timepoints
        for i_b in tindex_b:
            cc_b=cb[:,i_b]
            ti_b=round(t[i_b],5)
            plt.plot(x,cc_b,label='t={}'.format(ti_b))
        """
        
        #Logarthmic discreitzation of plotted timepoints
        lognt_b=np.log10(nt) #Logarthmic timepoints
        logspace_b=round((lognt_b)/tp_b,3) #Logarthmic timepoints
        logtindex_b=np.arange(0,lognt_b,logspace_b) #Logarthmic timepoints
        plt.figure(7*pc_i+1)
        for logi_b in logtindex_b:
            i_b=int(10**logi_b)
            cc_b=cb[:,i_b]
            ti_b=round(t[i_b],5)
            plt.plot(x,cc_b,label='t={}'.format(ti_b))
            
        plt.xlim(left=0,right=1)
        plt.ylim(bottom=0,top=upper_4)
        plt.xlabel('Position',fontsize=14)
        plt.ylabel('Dimensionless Concentration',fontsize=14)
        plt.title('Dimensionless Bound Concentration plot',fontsize=16)
        plt.xticks(fontsize=12)
        plt.yticks(fontsize=12)
        plt.legend(loc=(0.1,0.1))
        bound_filename_partial=f'Boundplot{pc_i}.png'
        bound_filename_full=os.path.join(internal_export_path,bound_filename_partial)
        plt.savefig(bound_filename_full)
        plt.close()
        pics_paragraph1=report.add_paragraph()
        pic1=pics_paragraph1.add_run()
        pic1.add_picture(unbound_filename_full, width=docx.shared.Inches(3))
        pic2=pics_paragraph1.add_run()
        pic2.add_picture(bound_filename_full, width=docx.shared.Inches(3))
        
        #Unbound NP Average Concetration Overtime
        plt.figure(7*pc_i+2)
        plt.plot(t,average_conc_overtime)
        plt.xlim(left=parameter_matrix[pc_i,2],right=parameter_matrix[pc_i,3])
        #plt.xlim(left=0,right=0.0005) #Manual Override of automatic x-axis limits
        plt.ylim(bottom=0,top=upper_2)
        plt.xlabel('Time',fontsize=14)
        plt.ylabel('Dimensionless Concentration',fontsize=14)
        plt.title('Average Dimensionless Unbound Concentration plot',fontsize=16)
        plt.xticks(fontsize=12)
        plt.yticks(fontsize=12)
        avgunbound_filename_partial=f'Avg Unboundplot{pc_i}.png'
        avgunbound_filename_full=os.path.join(internal_export_path,avgunbound_filename_partial)
        plt.savefig(avgunbound_filename_full)
        plt.close()
        
        #Unbound NP Change in Concentration vs Concentration
        plt.figure(7*pc_i+3)
        plt.plot(average_conc_overtime,change_in_concentration)
        plt.xlim(left=0,right=upper_2)
        plt.ylim(bottom=0,top=upper_3)
        plt.xlabel('Concentration',fontsize=14)
        plt.ylabel('Dimensionless Change in Concentration',fontsize=14)
        plt.title('Unbound dC vs C plot',fontsize=16)
        plt.xticks(fontsize=12)
        plt.yticks(fontsize=12)
        dCunbound_filename_partial=f'dC Unboundplot{pc_i}.png'
        dCunbound_filename_full=os.path.join(internal_export_path,dCunbound_filename_partial)
        plt.savefig(dCunbound_filename_full)
        plt.close()
        pics_paragraph2=report.add_paragraph()
        pic3=pics_paragraph2.add_run()
        pic3.add_picture(avgunbound_filename_full, width=docx.shared.Inches(3))
        pic4=pics_paragraph2.add_run()
        pic4.add_picture(dCunbound_filename_full, width=docx.shared.Inches(3))
        
        #Total NP Average Concetration Overtime
        plt.figure(7*pc_i+4)
        plt.plot(t,taverage_conc_overtime)
        plt.xlim(left=parameter_matrix[pc_i,2],right=parameter_matrix[pc_i,3])
        #plt.xlim(left=0,right=0.0005) #Manual Override of automatic x-axis limits
        plt.ylim(bottom=0,top=upper_5)
        plt.xlabel('Time',fontsize=14)
        plt.ylabel('Dimensionless Concentration',fontsize=14)
        plt.title('Average Dimensionless Total Concentration plot',fontsize=16)
        plt.xticks(fontsize=12)
        plt.yticks(fontsize=12)
        avgtotalNP_filename_partial=f'Avg totolNPplot{pc_i}.png'
        avgtotalNP_filename_full=os.path.join(internal_export_path,avgtotalNP_filename_partial)
        plt.savefig(avgtotalNP_filename_full)
        plt.close()
        
        #Total NP Change in Concentration vs Concentration
        plt.figure(7*pc_i+5)
        plt.plot(taverage_conc_overtime,tchange_in_concentration)
        plt.xlim(left=0,right=upper_5)
        plt.ylim(bottom=0,top=upper_6)
        plt.xlabel('Concentration',fontsize=14)
        plt.ylabel('Dimensionless Change in Concentration',fontsize=14)
        plt.title('Total dC vs C plot',fontsize=16)
        plt.xticks(fontsize=12)
        plt.yticks(fontsize=12)
        dCtotalNP_filename_partial=f'dC totolNPplot{pc_i}.png'
        dCtotalNP_filename_full=os.path.join(internal_export_path,dCtotalNP_filename_partial)
        plt.savefig(dCtotalNP_filename_full)
        plt.close()
        pics_paragraph3=report.add_paragraph()
        pic5=pics_paragraph3.add_run()
        pic5.add_picture(avgtotalNP_filename_full, width=docx.shared.Inches(3))
        pic6=pics_paragraph3.add_run()
        pic6.add_picture(dCtotalNP_filename_full, width=docx.shared.Inches(3))
        
        #Logarithms of Total NP Change in Concentration vs Concentration
        plt.figure(7*pc_i+6)
        plt.plot(logtavg_conc_overtime,logtchange_conc)
        plt.xlim(left=lower_7,right=upper_7)
        plt.ylim(bottom=lower_8,top=upper_8) 
        plt.xlabel('Log of Dimensionless Concentration',fontsize=14)
        plt.ylabel('Log of Dimensionless Change in Concentration',fontsize=14)
        plt.title('Total log(dC) vs log(C) plot',fontsize=16)
        plt.xticks(fontsize=12)
        plt.yticks(fontsize=12)
        logdCtotalNP_filename_partial=f'logdC totolNPplot{pc_i}.png'
        logdCtotalNP_filename_full=os.path.join(internal_export_path,logdCtotalNP_filename_partial)
        plt.savefig(logdCtotalNP_filename_full)
        plt.close()
        pics_paragraph4=report.add_paragraph()
        pic7=pics_paragraph4.add_run()
        pic7.add_picture(logdCtotalNP_filename_full, width=docx.shared.Inches(3))
        
        #Unbound Concentration Animation
        unbound_anim_fig=plt.figure()
        unbound_anim_plot=plt.plot([])
        unbound_anim_holder=unbound_anim_plot[0]
        plt.xlim(left=0,right=1)
        plt.ylim(bottom=0,top=upper_1)
        plt.xlabel('Position',fontsize=14)
        plt.ylabel('Dimensionless Concentration',fontsize=14)
        plt.title('Dimensionless Unbound Concentration',fontsize=16)
        plt.xticks(fontsize=12)
        plt.yticks(fontsize=12)
        tp_u_anim=100 #number of time points to plot for the animation for unbound concntration
        space_u_anim=int((nt-1)/tp_u_anim) #space between timepoints    
        def unbound_animate(frame):
            #update plot
            c_u_plot=cu[:,frame*space_u_anim]
            unbound_anim_holder.set_data((x,c_u_plot))
        unbound_anim=anim.FuncAnimation(unbound_anim_fig,unbound_animate,frames=tp_u_anim,interval=100)
        unbound_anim_filename_partial=f'unboun_anim{pc_i}.gif'
        unbound_anim_filename_full=os.path.join(internal_export_path,unbound_anim_filename_partial)
        unbound_anim.save(unbound_anim_filename_full)
        pic7.add_picture(unbound_anim_filename_full, width=docx.shared.Inches(3))
    return report


"""
Created on Sat Oct 31 18:13:53 2020

@author: joshuaprince

Purpose: Script to auto-generate report from data, including plotting of key figures and generation of an animated plot

Version 1.6

Changes from 1.5 to 1.6 (11/3/2020 5:25 pm):
    -going from linear discretization of time points to plot to logarthmic discretization

Differences between version 1.4 and 1.5 (11/3/2020 4:35 pm):
    -Added version number to script, along with reporting versions for each script

Differences between version 1.3 and 1.4 (11/3/2020 8:15 am):
    -Added plot close functionalities for all plots (python was complaining about holding so many plots in memory)
    -Lines 109, 132, 152, 167, 187, 202, 222: Added plt.close() 

Differences between 1.2 and 1.3 (11/3/2020 8:00 am)
-Changed file-path of animated unbound gif from direct export folder to internal export folder, since it is now incorporated into report
-Line 28: Got rid of direct_export_path from function call (no longer needed)
-Line 235: Changed "direct_export_path" to "internal_export_path"

Differences between 1.1 and 1.2
-Added lower bound calculations (needed for log plot), incorporated into log plotting
"""